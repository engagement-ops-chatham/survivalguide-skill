import argparse
import csv
import json
from pathlib import Path

from docx import Document


def render_guide(records: list[dict], output_docx):
    doc = Document()
    if records:
        doc.add_heading(records[0]["conference_name"], level=0)
    for record in records:
        doc.add_heading(record["source_firm_name"], level=1)
        doc.add_paragraph(f"Primary Contact: {', '.join(record['primary_contacts'])}")
        doc.add_paragraph(f"Contact Summary: {record['contact_summary']}")
        doc.add_paragraph(f"Company Summary: {record['company_summary']}")
        doc.add_paragraph(
            f"Relationship / Opportunity Context: {record['relationship_context']}"
        )
        doc.add_paragraph(record["review_line"])
        doc.add_paragraph("Sources")
        for source in record["sources"]:
            doc.add_paragraph(
                f'{source["label"]}: {source["url"]}', style="List Bullet"
            )
    doc.save(output_docx)


def write_qa_csv(records: list[dict], output_csv):
    fields = [
        "conference_name",
        "attendee_name",
        "source_firm_name",
        "hubspot_match_name",
        "hubspot_match_confidence",
        "hubspot_url",
        "linkedin_url",
        "review_flags",
    ]
    with Path(output_csv).open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        for record in records:
            row = {key: record.get(key, "") for key in fields}
            row["review_flags"] = ";".join(record.get("review_flags", []))
            writer.writerow(row)


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser()
    parser.add_argument("--records", required=True)
    parser.add_argument("--out-docx", required=True)
    parser.add_argument("--out-csv", required=True)
    return parser


def main() -> int:
    args = _build_parser().parse_args()
    records = json.loads(Path(args.records).read_text(encoding="utf-8"))
    render_guide(records, args.out_docx)
    write_qa_csv(records, args.out_csv)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
