import csv
import json
import shutil
import sys
import unittest
import uuid
from contextlib import contextmanager
from pathlib import Path
from unittest.mock import patch

from docx import Document

REPO_ROOT = Path(__file__).resolve().parents[1]
SCRATCH_ROOT = REPO_ROOT.parent / "test-scratch-render"
sys.path.insert(0, str(REPO_ROOT / "conference-survival-guide" / "scripts"))

from render_survival_guide import main, render_guide, write_qa_csv


@contextmanager
def make_workspace_tempdir():
    SCRATCH_ROOT.mkdir(exist_ok=True)
    tmp_path = SCRATCH_ROOT / f"tmp_test_{uuid.uuid4().hex}"
    tmp_path.mkdir()
    try:
        yield tmp_path
    finally:
        shutil.rmtree(tmp_path)


def sample_record() -> dict:
    return {
        "conference_name": "Industrials Heavy Hitters 2026",
        "attendee_name": "Daniel Troy",
        "source_firm_name": "Acacia",
        "primary_contacts": ["Daniel Troy"],
        "contact_summary": (
            "Daniel Troy is a principal at Acacia. "
            "He focuses on technology and industrial investments."
        ),
        "company_summary": (
            "Acacia is an operationally engaged investor. "
            "The firm backs technology and services businesses."
        ),
        "relationship_context": (
            "Prospect. No confirmed client match. One open opportunity exists."
        ),
        "review_line": (
            'Review: Source firm "Acacia" | HubSpot match "Acacia Group" | '
            "Confidence: likely | HubSpot: https://example.test | "
            "Source tier: LinkedIn | Flags: name-match-review"
        ),
        "sources": [
            {
                "label": "LinkedIn",
                "url": "https://linkedin.example/daniel-troy",
                "tier": "linkedin",
            },
            {
                "label": "Company Site",
                "url": "https://acacia.example/about",
                "tier": "company-site",
            },
        ],
        "hubspot_match_name": "Acacia Group",
        "hubspot_match_confidence": "likely",
        "hubspot_url": "https://example.test",
        "linkedin_url": "https://linkedin.example/daniel-troy",
        "review_flags": ["name-match-review"],
    }


class RenderGuideTest(unittest.TestCase):
    def test_render_guide_writes_required_sections(self):
        records = [sample_record()]
        with make_workspace_tempdir() as tmp_path:
            docx_path = tmp_path / "guide.docx"
            render_guide(records, docx_path)
            doc = Document(docx_path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertIn("Industrials Heavy Hitters 2026", text)
        self.assertIn("Daniel Troy", text)
        self.assertIn("Primary Contact: Daniel Troy", text)
        self.assertIn(
            "Contact Summary: Daniel Troy is a principal at Acacia. He focuses on technology and industrial investments.",
            text,
        )
        self.assertIn(
            "Company Summary: Acacia is an operationally engaged investor. The firm backs technology and services businesses.",
            text,
        )
        self.assertIn(
            "Relationship / Opportunity Context: Prospect. No confirmed client match. One open opportunity exists.",
            text,
        )
        self.assertIn("Review:", text)
        self.assertIn("Sources", text)
        self.assertIn("LinkedIn: https://linkedin.example/daniel-troy", text)

    def test_write_qa_csv_writes_expected_columns(self):
        records = [sample_record()]
        with make_workspace_tempdir() as tmp_path:
            csv_path = tmp_path / "guide.csv"
            write_qa_csv(records, csv_path)
            with csv_path.open("r", encoding="utf-8", newline="") as handle:
                reader = csv.DictReader(handle)
                rows = list(reader)
                headers = reader.fieldnames
        self.assertEqual(
            headers,
            [
                "conference_name",
                "attendee_name",
                "source_firm_name",
                "hubspot_match_name",
                "hubspot_match_confidence",
                "hubspot_url",
                "linkedin_url",
                "review_flags",
            ],
        )
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["attendee_name"], "Daniel Troy")
        self.assertEqual(rows[0]["hubspot_match_confidence"], "likely")
        self.assertEqual(rows[0]["hubspot_url"], "https://example.test")
        self.assertEqual(
            rows[0]["linkedin_url"], "https://linkedin.example/daniel-troy"
        )
        self.assertEqual(rows[0]["review_flags"], "name-match-review")

    def test_main_writes_docx_and_csv_from_cli_inputs(self):
        records = [sample_record()]
        with make_workspace_tempdir() as tmp_path:
            records_path = tmp_path / "records.json"
            docx_path = tmp_path / "guide.docx"
            csv_path = tmp_path / "guide.csv"
            records_path.write_text(json.dumps(records), encoding="utf-8")

            argv = [
                "render_survival_guide.py",
                "--records",
                str(records_path),
                "--out-docx",
                str(docx_path),
                "--out-csv",
                str(csv_path),
            ]
            with patch.object(sys, "argv", argv):
                exit_code = main()

            doc = Document(docx_path)
            with csv_path.open("r", encoding="utf-8", newline="") as handle:
                rows = list(csv.DictReader(handle))

        self.assertEqual(exit_code, 0)
        self.assertIn(
            "Daniel Troy", "\n".join(paragraph.text for paragraph in doc.paragraphs)
        )
        self.assertEqual(rows[0]["attendee_name"], "Daniel Troy")


if __name__ == "__main__":
    unittest.main()
